import io
import pandas as pd
from pypdf import PdfWriter, PdfReader
import docx
from markdown_it import MarkdownIt
from bs4 import BeautifulSoup
from xhtml2pdf import pisa

def _read_file_content(file):
    encodings_to_try = ['utf-8', 'latin-1', 'windows-1252']
    file.seek(0)
    for encoding in encodings_to_try:
        try:
            file.seek(0)
            return file.read().decode(encoding)
        except (UnicodeDecodeError, AttributeError):
            continue
    
    file.seek(0)
    return file.read().decode('utf-8', errors='ignore')

def merge_as_txt(files):
    output_string = io.StringIO()
    for file in files:
        file_extension = file.name.split('.')[-1].lower()
        content = ""
        try:
            if file_extension == 'pdf':
                file.seek(0)
                reader = PdfReader(file)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            elif file_extension == 'docx':
                file.seek(0)
                doc = docx.Document(file)
                for para in doc.paragraphs:
                    content += para.text + "\n"
            else: # For txt, md, html, csv etc.
                content = _read_file_content(file)

            output_string.write(content)
            output_string.write(f"\n\n--- End of File: {file.name} ---\n\n")
        except Exception as e:
            output_string.write(f"\n\n--- Could not process file: {file.name} (Error: {e}) ---\n\n")
    return output_string.getvalue().encode('utf-8')

def merge_as_csv(files):
    dataframes = []
    ignored_files = []
    for file in files:
        if file.name.endswith('.csv'):
            try:
                file.seek(0)
                content = _read_file_content(file)
                df = pd.read_csv(io.StringIO(content))
                dataframes.append(df)
            except Exception:
                ignored_files.append(file.name)
        else:
            ignored_files.append(file.name)
    
    if not dataframes:
        return None, ignored_files

    merged_df = pd.concat(dataframes, ignore_index=True)
    csv_buffer = io.StringIO()
    merged_df.to_csv(csv_buffer, index=False)
    return csv_buffer.getvalue().encode('utf-8'), ignored_files

def _convert_to_html_fragment(file):
    file_extension = file.name.split('.')[-1].lower()
    
    if file_extension == 'html':
        content = _read_file_content(file)
        soup = BeautifulSoup(content, 'lxml')
        return str(soup.body) if soup.body else ''
    elif file_extension == 'md':
        content = _read_file_content(file)
        md = MarkdownIt()
        return md.render(content)
    elif file_extension == 'csv':
        file.seek(0)
        content = _read_file_content(file)
        df = pd.read_csv(io.StringIO(content))
        return df.to_html(index=False)
    elif file_extension == 'txt':
        content = _read_file_content(file)
        return f'<pre>{content}</pre>'
    elif file_extension == 'docx':
        file.seek(0)
        doc = docx.Document(file)
        content = "\n".join([para.text for para in doc.paragraphs])
        return f'<div>{"<p>" + content.replace(chr(10), "</p><p>") + "</p>"}</div>'
    elif file_extension == 'pdf':
        file.seek(0)
        reader = PdfReader(file)
        content = "".join([page.extract_text() for page in reader.pages])
        return f'<div>{"<p>" + content.replace(chr(10), "</p><p>") + "</p>"}</div>'
    return ""

def merge_as_html(files):
    html_fragments = []
    for file in files:
        try:
            html_fragments.append(_convert_to_html_fragment(file))
        except Exception as e:
            html_fragments.append(f"<hr><p>Could not process file: {file.name} (Error: {e})</p>")

    full_html_content = "<hr>".join(html_fragments)
    
    final_html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>Merged Document</title>
        <style>
            body {{ font-family: sans-serif; margin: 2em; }}
            table {{ border-collapse: collapse; width: 100%; }}
            th, td {{ border: 1px solid #dddddd; text-align: left; padding: 8px; }}
            th {{ background-color: #f2f2f2; }}
            hr {{ margin: 2em 0; border: 1px solid #ccc; }}
            pre {{ background-color: #f5f5f5; padding: 1em; white-space: pre-wrap; word-wrap: break-word; }}
        </style>
    </head>
    <body>
        {full_html_content}
    </body>
    </html>
    """
    return final_html.encode('utf-8')

def merge_as_pdf(files):
    pdf_files = [f for f in files if f.name.endswith('.pdf')]
    other_files = [f for f in files if not f.name.endswith('.pdf')]
    
    writer = PdfWriter()
    
    for pdf_file in pdf_files:
        try:
            pdf_file.seek(0)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                writer.add_page(page)
        except Exception:
            # If a PDF is corrupted, we skip it but can convert its text
            other_files.append(pdf_file)

    if other_files:
        html_content_bytes = merge_as_html(other_files)
        
        pdf_output = io.BytesIO()
        pisa_status = pisa.CreatePDF(io.BytesIO(html_content_bytes), dest=pdf_output)
        
        if not pisa_status.err:
            pdf_output.seek(0)
            html_as_pdf_reader = PdfReader(pdf_output)
            for page in html_as_pdf_reader.pages:
                writer.add_page(page)

    pdf_bytes_io = io.BytesIO()
    writer.write(pdf_bytes_io)
    return pdf_bytes_io.getvalue()
